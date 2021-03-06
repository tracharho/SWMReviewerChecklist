from docx import Document
from docx.shared import Inches, Pt
import os
from datetime import datetime

def clearDocuments():
     foo = os.getcwd() + "/app/static"
     files_in_directory = os.listdir(foo)
     filtered_files = [file for file in files_in_directory if file.endswith(".docx")]
     for file in filtered_files:
          path_to_file = os.path.join(foo, file)
          os.remove(path_to_file)

def makeLetter(reviewername, recipientname, projectname, dscnumber, comments):
    picture_path = '/app/static/LetterHead.PNG'
    letter_path = '/app/static/letters'
    clearDocuments()
    doc = Document()
    sections = doc.sections
    sections[0].left_margin,sections[0].right_margin,sections[0].top_margin,sections[0].bottom_margin = Inches(0.75), Inches(0.75), Inches(0.5), Inches(0.5)    # Adding the City Logo
    doc.add_picture(os.getcwd()+picture_path)
    Para1 = doc.add_paragraph('INTER-OFFICE MEMORANDUM')
    Para1.runs[0].style = 'Title Char'
    Para1.add_run('\n\nDATE:').bold=True
    Para1.add_run('\t\t{}'.format(datetime.today().strftime('%m/%d/%Y')))
    Para1.add_run('\nTO:').bold=True
    Para1.add_run('\t\t{}'.format(recipientname))
    Para1.add_run('\nFROM:').bold = True
    Para1.add_run('\t{}'.format(reviewername))
    Para1.add_run('\nSUBJECT:').bold = True
    Para1.add_run('\t{}'.format(projectname))
    Para1.add_run('\n\t\tDSC Engineering Review Comments')
    Para1.add_run('\n\t\tDSC File: {}'.format(dscnumber)).add_break()

    Para2 = doc.add_paragraph('DSC cannot approve the project at this time. The following comments must first be addressed:\n')
    i = 1
    for comment in comments:
        Para2.add_run('\n{}:  {}'.format(i,comment))
        i += 1
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    for para in doc.paragraphs:
        para.style = doc.styles['Normal']
    
    letter_name = "Review Comments for DSC# {} {}.docx".format(dscnumber,projectname)
    str(os.path.abspath(os.getcwd()))
    letter_path = str(os.path.abspath(os.getcwd()))+ letter_path + letter_name
    doc.save(letter_path)
    return letter_path, letter_name
